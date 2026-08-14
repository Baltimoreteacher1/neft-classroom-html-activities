"""Emit SY26-27 Grade 6 Math Year Planning Calendar.docx — fully editable Word."""

import json, datetime as dt, calendar as cal, collections
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from plan_core import *

OUT = "/Users/joelneft/Desktop/SY26-27 Grade 6 Math Planning Calendar/SY26-27 Grade 6 Math Year Planning Calendar.docx"
P = json.load(open("plan.json"))
rows = P["rows"]
for r in rows:
    r["d"] = dt.date.fromisoformat(r["date"])
BYDATE = {r["d"]: r for r in rows}
INK, MUTED, RULE = "1F2A37", "5A6472", "C9CFD8"


# ------------------------------------------------------------------ helpers
def shade(cell, hexcolor):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def borders(
    obj,
    color=RULE,
    sz=4,
    edges=("top", "left", "bottom", "right", "insideH", "insideV"),
):
    pr = obj._tc.get_or_add_tcPr() if hasattr(obj, "_tc") else obj._tbl.tblPr
    tag = "w:tcBorders" if hasattr(obj, "_tc") else "w:tblBorders"
    b = OxmlElement(tag)
    for e in edges:
        el = OxmlElement(f"w:{e}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        b.append(el)
    pr.append(b)


def cell_margins(table, top=60, bottom=60, left=90, right=90):
    m = OxmlElement("w:tblCellMar")
    for k, v in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement(f"w:{k}")
        e.set(qn("w:w"), str(v))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    table._tbl.tblPr.append(m)


def keep_with_next(par, on=True):
    par.paragraph_format.keep_with_next = on


def repeat_header(row):
    tr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr.append(el)


def cant_split(row):
    tr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:cantSplit")
    el.set(qn("w:val"), "true")
    tr.append(el)


def txt(
    cell, text, size=9, bold=False, color=INK, italic=False, space_after=0, first=False
):
    p = (
        cell.paragraphs[0]
        if first and not cell.paragraphs[0].text
        else cell.add_paragraph()
    )
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def hyperlink(par, url, label, size=8):
    part = par.part
    rid = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), rid)
    r = OxmlElement("w:r")
    pr = OxmlElement("w:rPr")
    for tag, val in (("w:color", "1F4E79"), ("w:sz", str(int(size * 2)))):
        e = OxmlElement(tag)
        e.set(qn("w:val"), val.replace("w:", ""))
        pr.append(e)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    pr.append(u)
    r.append(pr)
    t = OxmlElement("w:t")
    t.text = label
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    h.append(r)
    par._p.append(h)


def H(
    doc,
    text,
    level=1,
    size=None,
    color="1F3864",
    space_before=14,
    space_after=6,
    page_break=False,
):
    p = doc.add_paragraph()
    if page_break:
        p.add_run().add_break(WD_BREAK.PAGE)
    p.style = doc.styles[f"Heading {level}"]
    r = p.add_run(text)
    r.font.size = Pt(size or {1: 18, 2: 14, 3: 11.5}[level])
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = True
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    return p


def body(doc, text, size=10, italic=False, color=INK, after=6, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.italic = italic
    r.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(after)
    return p


def mktable(doc, headers, widths, accent="1F3864", size=9):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    borders(t)
    cell_margins(t)
    for i, (h, w) in enumerate(zip(headers, widths)):
        c = t.rows[0].cells[i]
        c.width = Inches(w)
        shade(c, accent)
        txt(c, h, size=size, bold=True, color="FFFFFF", first=True)
    repeat_header(t.rows[0])
    cant_split(t.rows[0])
    return t


def addrow(t, widths, vals, size=9, fills=None, bolds=None):
    r = t.add_row()
    cant_split(r)
    for i, v in enumerate(vals):
        c = r.cells[i]
        c.width = Inches(widths[i])
        if fills and fills[i]:
            shade(c, fills[i])
        txt(c, str(v), size=size, bold=bool(bolds and bolds[i]), first=True)
    return r


def accent_of(u):
    return UNIT_ACCENT.get(u, "555555")


# ------------------------------------------------------------------ document
doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = Inches(11), Inches(8.5)
for a in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
    setattr(sec, a, Inches(0.5))
ft = sec.footer.paragraphs[0]
ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = ft.add_run(
    "SY26-27 Grade 6 Math — Year Planning Calendar   ·   editable Word document   ·   page "
)
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor.from_string(MUTED)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
ft._p.append(fld)

# ============================================================ TITLE
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("Grade 6 Mathematics")
r.font.size = Pt(30)
r.bold = True
r.font.color.rgb = RGBColor.from_string("1F3864")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Year Planning Calendar — School Year 2026–2027")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor.from_string(MUTED)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(18)
r = p.add_run("Course 1  ·  August 24, 2026 – June 11, 2027")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(INK)
for line in [
    "Built from the SY26-27 school calendar and the Course 1 scope & sequence.",
    "Lesson titles, standards, objectives and links come from the current EduWonderLab curriculum manifest.",
    "Every date, lesson, note and table cell in this document is native, editable Word text.",
]:
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = q.add_run(line)
    rr.font.size = Pt(9.5)
    rr.italic = True
    rr.font.color.rgb = RGBColor.from_string(MUTED)
    q.paragraph_format.space_after = Pt(2)

# ============================================================ 1. YEAR AT A GLANCE
H(doc, "Section 1 · Year at a Glance", 1, page_break=True, space_before=0)
body(
    doc,
    "Units appear in the order the district scope & sequence teaches them — which is not "
    "numerical order. Unit numbers are the district's; EduWonderLab uses the same numbering.",
    size=9.5,
    italic=True,
    color=MUTED,
)

W = [0.85, 3.15, 1.0, 1.0, 0.85, 0.8, 0.85, 1.55]
t = mktable(
    doc,
    [
        "Unit",
        "Unit Name",
        "Approx. Start",
        "Approx. End",
        "Instr. Days",
        "Lessons",
        "Assessment",
        "Culminating Project",
    ],
    W,
)
lessons, sg, cu, eou = load_ewl()
for u in P["units"]:
    ur = [r for r in rows if r["unit_key"] == u["key"]]
    nl = len(
        {
            r["lesson_id"]
            for r in ur
            if r["day_type"] == "Core Lesson" and r["lesson_id"]
        }
    )
    npj = sum(1 for r in ur if r["day_type"] == "Project")
    a = accent_of(u["ewl"])
    row = addrow(
        t,
        W,
        [
            (f"Unit {u['ewl']}" if u["ewl"] else "MSTAR"),
            u["label"].split(": ", 1)[-1],
            dt.date.fromisoformat(u["start"]).strftime("%b %-d, %Y"),
            dt.date.fromisoformat(u["end"]).strftime("%b %-d, %Y"),
            len(ur),
            nl or "—",
            u["marker"] or "—",
            f"{npj} days" if npj else "—",
        ],
    )
    shade(row.cells[0], a)
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("FFFFFF")
    row.cells[0].paragraphs[0].runs[0].bold = True

H(doc, "Quarters, breaks and the testing window", 2)
W2 = [2.2, 2.3, 5.5]
t = mktable(doc, ["Marker", "Dates", "Planning impact"], W2, accent="44546A")
inst_by_q = collections.Counter(
    r["quarter"] for r in rows if r["day_type"] != "No Instruction"
)
for q, a, b in QUARTERS:
    addrow(
        t,
        W2,
        [
            q,
            f"{a.strftime('%b %-d, %Y')} – {b.strftime('%b %-d, %Y')}",
            f"{inst_by_q[q]} instructional dates. Units are not forced to end on the quarter "
            f"boundary; several deliberately cross it.",
        ],
    )
for n, a, b in BREAKS:
    addrow(
        t,
        W2,
        [
            n,
            f"{a.strftime('%b %-d')} – {b.strftime('%b %-d, %Y')}",
            "No multi-day lesson straddles this break; the day after is a retrieval / catch-up day.",
        ],
    )
addrow(
    t,
    W2,
    [
        "Thanksgiving",
        "Nov 25 (Wellness) – Nov 27, 2026",
        "Three consecutive non-student days mid-week. Nov 24 is the last teaching day.",
    ],
)
addrow(
    t,
    W2,
    [
        "State testing window",
        "Mar 29 – May 28, 2027",
        "Labelled 'MCAP Window' on the calendar; the scope & sequence calls the test MSTAR. "
        "Exact in-building dates are NOT in either source — confirm locally and re-anchor the "
        "MSTAR block on the Year Pacing sheet.",
    ],
)
addrow(
    t,
    W2,
    [
        "MSTAR block",
        "May 18 – May 28, 2027",
        "3 preparation days + 6 administration days, placed where the scope & sequence places them.",
    ],
)
addrow(
    t,
    W2,
    [
        "Shortened days",
        "Oct 15, Oct 23, Jan 15, Mar 25, Jun 11",
        "No new content launches on these days — they carry continuation, catch-up, project or showcase work.",
    ],
)

H(doc, "Where the flex lives", 2)
flexrows = [r for r in rows if r["day_type"] in ("Flex", "Catch-Up", "Review")]
body(
    doc,
    f"{len(flexrows)} of the {sum(1 for r in rows if r['day_type'] != 'No Instruction')} "
    f"instructional days are deliberately uncommitted or repair-focused: "
    f"{sum(1 for r in flexrows if r['day_type'] == 'Flex')} open Flex days, "
    f"{sum(1 for r in flexrows if r['day_type'] == 'Catch-Up')} Catch-Up station days anchored to "
    f"specific lessons, and {sum(1 for r in flexrows if r['day_type'] == 'Review')} unit Review days. "
    f"Every one is listed on the Flex Days sheet of the pacing workbook.",
    size=10,
)

# ============================================================ 2. MONTHLY CALENDARS
H(doc, "Section 2 · Monthly Calendars", 1, page_break=True, space_before=0)
body(
    doc,
    "One page per month, Monday–Friday. Each cell shows the date, the unit and lesson, and a "
    "short day label. Type directly into any cell.",
    size=9.5,
    italic=True,
    color=MUTED,
)

months = []
d = dt.date(2026, 8, 1)
while d <= dt.date(2027, 6, 1):
    months.append((d.year, d.month))
    d = (d.replace(day=28) + dt.timedelta(days=7)).replace(day=1)

DAYLBL = {
    "Core Lesson": "",
    "Continued Lesson": "continued",
    "Assessment": "ASSESSMENT",
    "Project": "PROJECT",
    "Review": "Review",
    "Flex": "FLEX",
    "Catch-Up": "Catch-Up",
    "MCAP / Testing": "TESTING",
}

for mi, (y, m) in enumerate(months):
    H(
        doc,
        f"{cal.month_name[m]} {y}",
        2,
        size=16,
        page_break=(mi > 0),
        space_before=0 if mi else 8,
    )
    WC = [2.0] * 5
    t = mktable(
        doc,
        ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"],
        WC,
        accent="44546A",
        size=9.5,
    )
    first = dt.date(y, m, 1)
    start = first - dt.timedelta(days=first.weekday())
    last = dt.date(y, m, cal.monthrange(y, m)[1])
    week = start
    while week <= last:
        days5 = [week + dt.timedelta(days=i) for i in range(5)]
        if not any(
            d.month == m and (d in BYDATE or dt.date(2026, 8, 18) <= d <= dt.date(2026, 8, 21))
            for d in days5
        ):
            week += dt.timedelta(days=7)
            continue
        r = t.add_row()
        cant_split(r)
        r.height = Inches(1.02)
        for i in range(5):
            day = week + dt.timedelta(days=i)
            c = r.cells[i]
            c.width = Inches(2.0)
            if day.month != m:
                shade(c, "FAFAFA")
                week_ = None
                continue
            rec = BYDATE.get(day)
            hd = txt(
                c, day.strftime("%b %-d"), size=8.5, bold=True, color=MUTED, first=True
            )
            if rec is None:
                if dt.date(2026, 8, 18) <= day <= dt.date(2026, 8, 21):
                    shade(c, "F2F2F2")
                    txt(c, "School & Systemic PD", size=8, color=MUTED, italic=True)
                else:
                    shade(c, "FAFAFA")
                continue
            if rec["day_type"] == "No Instruction":
                shade(c, "F0F0F0")
                txt(
                    c,
                    rec["cal_note"].split(";")[0] or rec["status"],
                    size=8,
                    color=MUTED,
                    italic=True,
                )
                continue
            if rec["minutes"] == "Half Day":
                shade(c, "FFF6E5")
                txt(c, "½ DAY / EARLY RELEASE", size=7.5, bold=True, color="8A5A00")
            unit = f"U{rec['ewl_unit']}" if rec["ewl_unit"] else "MSTAR"
            lid = rec["lesson_id"].replace("-", ".") if rec["lesson_id"] else ""
            head = (
                f"{unit} · L{lid}"
                if lid and rec["day_type"] in ("Core Lesson", "Continued Lesson")
                else (f"{unit} · {DAYLBL.get(rec['day_type']) or rec['day_type']}")
            )
            txt(c, head, size=8.5, bold=True, color=accent_of(rec["ewl_unit"]))
            title = rec["title"]
            if len(title) > 84:
                title = title[:81].rstrip() + "…"
            txt(c, title, size=8, color=INK)
            tail = []
            if rec["day_type"] == "Continued Lesson":
                tail.append("continued")
            for note in rec["cal_note"].split(";"):
                note = note.strip()
                if note and (
                    "Q" in note
                    and "STARTS" in note.upper()
                    or "progress-report" in note
                    or "FIRST DAY" in note
                    or "LAST DAY" in note
                ):
                    tail.append(note)
            ded = []
            for x in tail:
                nx = x.lower().replace(" ", "")
                if any(nx in y.lower().replace(" ", "") or
                       y.lower().replace(" ", "") in nx for y in ded):
                    continue
                ded.append(x)
            tail = ded
            if tail:
                txt(
                    c,
                    " · ".join(tail),
                    size=7.5,
                    italic=True,
                    color="8A5A00" if rec["minutes"] == "Half Day" else MUTED,
                )
        week += dt.timedelta(days=7)
    # month legend
    ev = [
        f"{d.strftime('%-d')}: {CLOSURES[d][1]}"
        for d in sorted(CLOSURES)
        if d.year == y and d.month == m
    ]
    ev += [
        f"{d.strftime('%-d')}: {HALF_DAYS[d]}"
        for d in sorted(HALF_DAYS)
        if d.year == y and d.month == m
    ]
    for d in sorted(MILESTONES):
        if d.year != y or d.month != m:
            continue
        key = MILESTONES[d].lower().replace(" ", "")
        if any(e.startswith(f"{d.strftime('%-d')}:") and key[:8] in
               e.lower().replace(" ", "") for e in ev):
            continue  # already covered by the half-day / closure label for this date
        ev.append(f"{d.strftime('%-d')}: {MILESTONES[d]}")
    if ev:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        rr = p.add_run("Calendar notes:  " + "   ·   ".join(dict.fromkeys(ev)))
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor.from_string(MUTED)

# ============================================================ 3. WEEKLY PAGES
H(doc, "Section 3 · Weekly Planning Pages", 1, page_break=True, space_before=0)
body(
    doc,
    "One page per instructional week. The Whole-Group and Small-Group columns carry live links "
    "to the EduWonderLab lesson and its two small-group versions. Every cell is editable — the "
    "Differentiation and Assessment/Notes columns are intentionally left light for your own planning.",
    size=9.5,
    italic=True,
    color=MUTED,
)

byweek = collections.OrderedDict()
for r in rows:
    if r["day_type"] != "No Instruction" or True:
        byweek.setdefault(r["week"], []).append(r)

WW = [0.62, 0.68, 2.15, 2.35, 2.35, 1.85]
FIRST_WEEK = next(w for w, rr in byweek.items() if any(x["day_type"] != "No Instruction" for x in rr))
FIRST_UNIT = P["units"][0]["key"]
for wk, wrows in byweek.items():
    if not any(r["day_type"] != "No Instruction" for r in wrows):
        continue
    a, b = wrows[0]["d"], wrows[-1]["d"]
    span = f"{a.strftime('%B %-d')}–{b.strftime('%-d') if a.month == b.month else b.strftime('%B %-d')}, {b.year}"
    H(doc, f"Week {wk} — {span}", 2, size=15,
      page_break=(wk != FIRST_WEEK), space_before=0 if wk != FIRST_WEEK else 10)
    units = list(dict.fromkeys(r["unit_label"] for r in wrows if r["unit_label"]))
    qs = list(dict.fromkeys(r["quarter"] for r in wrows if r["quarter"]))
    sub = "   |   ".join(units) + (f"      ·      {', '.join(qs)}" if qs else "")
    p = body(doc, sub, size=10.5, bold=True, after=2)
    p.runs[0].font.color.rgb = RGBColor.from_string(
        accent_of(next((r["ewl_unit"] for r in wrows if r["ewl_unit"]), None))
    )
    ndays = sum(1 for r in wrows if r["day_type"] != "No Instruction")
    val = sum(
        0.5 if r["minutes"] == "Half Day" else 1
        for r in wrows
        if r["day_type"] != "No Instruction"
    )
    body(
        doc,
        f"{ndays} instructional dates this week  ·  {val:g} instructional days",
        size=8.5,
        italic=True,
        color=MUTED,
        after=5,
    )

    t = mktable(
        doc,
        [
            "Day",
            "Lesson",
            "Objective / Standard",
            "Whole-Group Focus",
            "Small-Group / Differentiation",
            "Assessment / Notes",
        ],
        WW,
        size=8.5,
    )
    cell_margins(t, top=32, bottom=32, left=80, right=80)
    for r in wrows:
        row = t.add_row()
        cant_split(row)
        for i, w in enumerate(WW):
            row.cells[i].width = Inches(w)
        c0 = row.cells[0]
        txt(c0, r["d"].strftime("%a"), size=9, bold=True, first=True)
        txt(c0, r["d"].strftime("%-m/%-d"), size=8, color=MUTED)
        if r["day_type"] == "No Instruction":
            for i in range(6):
                shade(row.cells[i], "F0F0F0")
            txt(row.cells[1], r["status"], size=9, bold=True, color=MUTED, first=True)
            txt(
                row.cells[2],
                r["cal_note"],
                size=8.5,
                italic=True,
                color=MUTED,
                first=True,
            )
            continue
        if r["minutes"] == "Half Day":
            for i in range(6):
                shade(row.cells[i], "FFF6E5")
        # Lesson
        c1 = row.cells[1]
        lab = (
            f"L{r['lesson_id'].replace('-', '.')}"
            if r["lesson_id"] and r["day_type"] in ("Core Lesson", "Continued Lesson")
            else r["day_type"]
        )
        txt(c1, lab, size=9, bold=True, color=accent_of(r["ewl_unit"]), first=True)
        if r["day_type"] == "Continued Lesson":
            txt(c1, "continued", size=7.5, italic=True, color=MUTED)
        if r["minutes"] == "Half Day":
            txt(c1, "½ day", size=7.5, bold=True, color="8A5A00")
        # Objective / standard
        c2 = row.cells[2]
        if r["standard"]:
            txt(c2, r["standard"], size=8, bold=True, color=MUTED, first=True)
        txt(c2, r["objective"] or r["note"] or "—", size=7.8, first=not r["standard"])
        # Whole-group focus
        c3 = row.cells[3]
        txt(c3, r["title"], size=8, bold=True, first=True)
        if r["lesson_url"]:
            pp = c3.add_paragraph()
            pp.paragraph_format.space_after = Pt(0)
            hyperlink(pp, r["lesson_url"], "Interactive Whole Group Lesson", size=7.5)
        # Small group
        c4 = row.cells[4]
        if r["sg_urls"]:
            for i, u in enumerate(r["sg_urls"].split(" | "), 1):
                pp = c4.paragraphs[0] if i == 1 else c4.add_paragraph()
                pp.paragraph_format.space_after = Pt(0)
                hyperlink(pp, u, f"Small-Group Lesson · Group {i}", size=7.5)
        else:
            txt(c4, "", size=8.5, first=True)
        txt(row.cells[5], "", size=8.5, first=True)

    # notes areas
    H(doc, "Flex / Adjustment Notes", 3, size=10.5, space_before=7, space_after=3)
    nt = doc.add_table(rows=1, cols=1)
    borders(nt)
    cell_margins(nt, top=50, bottom=50)
    nt.rows[0].cells[0].width = Inches(10.0)
    nt.rows[0].height = Inches(0.42)
    txt(nt.rows[0].cells[0], "", size=9, first=True)

    H(doc, "If We Lose a Day", 3, size=10.5, space_before=7, space_after=3)
    soft = [r for r in wrows if r["day_type"] in ("Flex", "Catch-Up", "Review")]
    multi = [r for r in wrows if r["day_type"] == "Continued Lesson"]
    if soft:
        rec = (
            f"Absorb it with the {soft[0]['day_type']} day on {soft[0]['d'].strftime('%A %-m/%-d')} — "
            f"run it as a short retrieval opener instead of a full day."
        )
    elif multi:
        rec = (
            f"Continue {multi[0]['title'].split(' — ')[0]} into the next day and shift the rest of the "
            f"week forward one slot; the next Flex or Catch-Up day absorbs the shift."
        )
    else:
        rec = (
            "Shift the whole week forward one slot and absorb it at the next Catch-Up or Flex day. "
            "Do not drop a lesson — the unit assessment assumes all of them."
        )
    rt = doc.add_table(rows=1, cols=1)
    borders(rt)
    cell_margins(rt, top=45, bottom=45)
    rt.rows[0].cells[0].width = Inches(10.0)
    shade(rt.rows[0].cells[0], "F7F9FC")
    txt(rt.rows[0].cells[0], rec, size=8.5, italic=True, first=True)

# ============================================================ 4. UNIT OVERVIEWS
H(doc, "Section 4 · Unit Pacing Overviews", 1, page_break=True, space_before=0)
body(
    doc,
    "One page per unit, in teaching order. Use these when the year changes and a unit has to move.",
    size=9.5,
    italic=True,
    color=MUTED,
)

UW = [0.75, 3.5, 1.9, 0.6, 3.25]
for u in P["units"]:
    ur = [r for r in rows if r["unit_key"] == u["key"]]
    a = accent_of(u["ewl"])
    name = u["label"]
    H(doc, name, 2, size=15, color=a,
      page_break=(u["key"] != FIRST_UNIT), space_before=0 if u["key"] != FIRST_UNIT else 10)
    stds = sorted({r["standard"] for r in ur if r["standard"]})
    npj = [r for r in ur if r["day_type"] == "Project"]
    nas = [r for r in ur if r["day_type"] == "Assessment"]
    nfx = [r for r in ur if r["day_type"] in ("Flex", "Catch-Up", "Review")]
    MW = [1.75, 8.25]
    t = mktable(doc, ["Field", "Detail"], MW, accent=a, size=8.5)
    cell_margins(t, top=26, bottom=26, left=80, right=80)

    def mrow(k, v):
        addrow(t, MW, [k, v], size=8.6, bolds=[True, False])

    mrow(
        "Approximate dates",
        f"{dt.date.fromisoformat(u['start']).strftime('%B %-d, %Y')} – "
        f"{dt.date.fromisoformat(u['end']).strftime('%B %-d, %Y')}   "
        f"({len(ur)} instructional dates)",
    )
    mrow(
        "Scope & sequence budget",
        f"{u['budget']:g} instructional days"
        + (
            f"   ·   adjusted by {BUDGET_ADJUST[u['key']]:+g} day(s) — see Planning Notes"
            if u["key"] in BUDGET_ADJUST
            else ""
        ),
    )
    mrow("Standards", ", ".join(stds) if stds else "—")
    mrow(
        "Lessons",
        f"{len({r['lesson_id'] for r in ur if r['lesson_id'] and r['day_type'] == 'Core Lesson'})} lessons "
        f"across {sum(1 for r in ur if r['day_type'] in ('Core Lesson', 'Continued Lesson'))} days",
    )
    mrow(
        "Assessment",
        (
            f"{u['marker']} on {nas[0]['d'].strftime('%A, %B %-d')}"
            if nas
            else "— none in the scope & sequence"
        ),
    )
    mrow(
        "Culminating project",
        (
            f"{npj[0]['title'].split(' — ')[0]} · {len(npj)} days "
            f"({npj[0]['d'].strftime('%b %-d')} – {npj[-1]['d'].strftime('%b %-d')})"
        )
        if npj
        else "—",
    )
    mrow(
        "Built-in flex / reteach",
        f"{len(nfx)} day{'s' if len(nfx) != 1 else ''}: "
        + ", ".join(f"{r['d'].strftime('%b %-d')} ({r['day_type']})" for r in nfx)
        if nfx
        else "none",
    )
    interrupts = [
        f"{r['d'].strftime('%b %-d')} — {r['cal_note'].split(';')[0]}"
        for r in rows
        if u["start"] <= r["date"] <= u["end"] and r["day_type"] == "No Instruction"
    ]
    interrupts += [
        f"{r['d'].strftime('%b %-d')} — shortened day"
        for r in ur
        if r["minutes"] == "Half Day"
    ]
    mrow("Calendar interruptions", "; ".join(interrupts) if interrupts else "none")

    H(doc, "Lesson map", 3, size=10.5, space_before=7, space_after=3)
    t = mktable(
        doc, ["Lesson", "Title", "Planned Date(s)", "Days", "Notes"], UW, accent=a
    )
    grouped = []  # consecutive runs: multi-day lessons and multi-day projects group
    for r in ur:
        k = (r["lesson_id"] or r["title"].split(" — Day")[0],
             "L" if r["lesson_id"] else r["day_type"])
        if grouped and grouped[-1][0] == k:
            grouped[-1][1].append(r)
        else:
            grouped.append((k, [r]))
    for k, g in grouped:
        r0 = g[0]
        lab = (
            f"L{r0['lesson_id'].replace('-', '.')}"
            if r0["lesson_id"] and r0["day_type"] in ("Core Lesson", "Continued Lesson")
            else r0["day_type"]
        )
        dates = (
            f"{g[0]['d'].strftime('%b %-d')} – {g[-1]['d'].strftime('%b %-d')}"
            if len(g) > 1
            else g[0]["d"].strftime("%b %-d")
        )
        note = r0["note"] or ""
        if any(x["minutes"] == "Half Day" for x in g):
            note = (note + " · " if note else "") + "includes a shortened day"
        addrow(
            t, UW, [lab, r0["title"].split(" — Day")[0], dates, len(g), note], size=8.4
        )

# ============================================================ 5. ADJUSTMENT LOG
H(doc, "Section 5 · Pacing Adjustment Log", 1, page_break=True, space_before=0)
body(
    doc,
    "Record every change you make to the plan during the year. Keeping this current is what "
    "makes the calendar a planning system rather than a snapshot — and it is what you will hand "
    "the team when the sequence has to move.",
    size=10,
)
LW = [1.3, 2.6, 2.6, 1.9, 1.6]
t = mktable(
    doc,
    ["Date Changed", "Original Plan", "New Plan", "Reason", "Downstream Impact"],
    LW,
    accent="44546A",
)
addrow(t, LW, ["", "", "", "", ""], size=9)
addrow(t, LW, ["", "", "", "", ""], size=9)
for _ in range(14):
    r = t.add_row()
    cant_split(r)
    for i, w in enumerate(LW):
        r.cells[i].width = Inches(w)
        txt(r.cells[i], "", size=9, first=True)
    r.height = Inches(0.32)

H(doc, "Sources and assumptions", 2, page_break=True, space_before=0)
for line in [
    "Dates, closures, PD days, quarter starts, conference windows, early-release days and the testing "
    "window are taken from the SY26-27 school calendar. That calendar is authoritative wherever it "
    "disagrees with any other source.",
    "Unit order, unit names and per-unit day budgets are taken from the 'Course 1' sheet of the Course 1 "
    "scope & sequence. The teaching order there is not numerical — it runs Pre-Unit, 3, 4, 6, 7, 8, 9, 5, "
    "2, MSTAR, 10 — and this calendar preserves it exactly.",
    "Lesson IDs, titles, standards, objectives and all links are read from the current EduWonderLab "
    "curriculum manifest. No lesson title or URL in this document was invented.",
    "Two weeks in the scope & sequence assume more instructional days than the calendar actually has "
    "(Nov 2–6 and Mar 22–26), a total of 2.5 days. Those days were removed from flex in the units where "
    "they occur, not from lessons. The full reasoning is in 'SY26-27 Grade 6 Math Planning Notes.md'.",
    "Exact in-building MSTAR testing dates are not stated in either source and were not invented. The "
    "MSTAR block sits where the scope & sequence places it, inside the published window.",
    "Every planning decision that is an inference rather than a source fact is labelled in the Planning "
    "Notes file as INFERRED PLANNING DECISION or REQUIRES TEACHER CONFIRMATION.",
]:
    p = doc.add_paragraph(line, style="List Bullet")
    p.runs[0].font.size = Pt(10)
    p.paragraph_format.space_after = Pt(7)

doc.save(OUT)
print("wrote", OUT)
